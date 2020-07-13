from textfile_solution import PROJECT_PATH
from docx import Document
from PIL import Image, ImageDraw
from io import BytesIO
import psutil
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn

def create_doc():
    document = Document()
    document.add_paragraph('Hello,Word!')
    filename = PROJECT_PATH + '/file_location/demo.docx'
    document.save(filename)


def insert_graph_doc():
    document = Document()                        #新建文档 
    p = document.add_paragraph()                 #添加一个段落 
    r = p.add_run()                              #添加一个游程 
    img_size = 20
    for x in range(20):
        im = Image.new('RGB', (img_size, img_size), 'white')
        draw_obj = ImageDraw.Draw(im)
        draw_obj.ellipse((0, 0, img_size-1, img_size-1), fill=255-x)#画圆 
        fake_buf_file = BytesIO()          #用 BytesIO 将图片保存在内存里，减少磁盘操作 
        im.save(fake_buf_file, 'png')
        r.add_picture(fake_buf_file)            #在当前游程中插入图片 
        fake_buf_file.close()
        filename = PROJECT_PATH + '/file_location/demo.docx'
        document.save(filename)


def create_structure_doc():
    doc = Document()
    filename = PROJECT_PATH + '/file_location/demo-1.docx'
    doc.add_paragraph(u'Python 为什么这么受欢迎？', 'Title')
    doc.add_paragraph(u'作者','Subtitle')
    doc.add_paragraph(u'摘要：本文阐明了 Python 的优势...','Body Text 2')
    doc.add_paragraph(u'简单', 'Heading 1')
    doc.add_paragraph(u'易学')
    doc.add_paragraph(u'易用', 'Heading 2')
    doc.add_paragraph(u'功能强')
    p = doc.add_paragraph(u'贴合小年轻')
    p.style = 'Heading 2'
    doc.save(filename)


def style_doc():
    doc = Document()
    styles = doc.styles
    print("\n".join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))


def read_doc_test():
    from docx import Document
    import sys
    path = PROJECT_PATH + "/file_location/demo.docx"
    document = Document(path)
    for p in document.paragraphs:
        print(len(p.text))
        print(p.style.name)
        print(p.text)


def create_cell_doc_test():
    #获取当前计算机的配置 
    vmem = psutil.virtual_memory()
    vmem_dict = vmem._asdict()

    trow = 2
    tcol = len(vmem_dict.keys())
    #产生表格 

    document = Document()
    path = PROJECT_PATH + "/file_location/table.docx"
    table = document.add_table(rows=trow,cols=tcol,style = 'Table Grid')
    for col, info in enumerate(vmem_dict.keys()):
        table.cell(0,col).text = info
        if info == 'percent':
            table.cell(1,col).text = str(vmem_dict[info])+ '%'
        else:
            table.cell(1,col).text = str(vmem_dict[info]/(1024*1024)) + 'MB'
    document.save(path)


def operate_cell_doc_test():
    document = Document()
    table = document.add_table(rows=9, cols=10, style = 'Table Grid')
    cell_1 = table.cell(1,2)
    cell_2 = table.cell(4,6)
    cell_1.merge(cell_2)
    path = PROJECT_PATH + "/file_location/table-1.docx"
    path_2 = PROJECT_PATH + "/file_location/table-2.docx"
    document.save(path)

    document = Document(path)
    table = document.tables[0]
    for row,obj_row in enumerate(table.rows):
        for col,cell in enumerate(obj_row.cells):
            cell.text = cell.text + "%d,%d " % (row, col)
    document.save(path_2)


def operate_cell_doc_1_test():
    document = Document()
    path = PROJECT_PATH + "/file_location/table-step.docx"
    for row in range(9):
        t = document.add_table(rows=1, cols=1, style='Table Grid')
        t.autofit = False  # 很重要，必须设置 
        w = float(row) / 2.0
        t.columns[0].width = Inches(w)

    document.save(path)


def usage_of_style_doc():
    document = Document()
    styles = document.styles
    table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
    for style in table_styles:
        print(style.name)


    # 准备一些数据
    document = Document()
    vmem = psutil.virtual_memory()
    vmem_dict = vmem._asdict()
    trow = 2
    tcol = len(vmem_dict.keys())
    # 指定样式修饰表格
    table = document.add_table(rows=trow, cols=tcol, style='Colorful Grid Accent 4')
    for col, info in enumerate(vmem_dict.keys()):
        table.cell(0, col).text = info
        if info == 'percent':
            table.cell(1, col).text = str(vmem_dict[info]) + '%'
        else:
            table.cell(1, col).text = str(vmem_dict[info] / (1024 * 1024)) + 'M'
    document.save(PROJECT_PATH + "/file_location/table-style.docx")

    # 创建样式和设置字体
    doc = Document()
    for i in range(10):
        p = doc.add_paragraph(u'段落 % d' % i)
        style = doc.styles.add_style('UserStyle%d' % i, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(i + 20)
        p.style = style
    doc.save(PROJECT_PATH + '/file_location/style-1.docx')

    # 设置字体样式
    doc = Document()
    p = doc.add_paragraph()
    text_str = u'好好学习Python，努力做开发专家，成为最牛的程序员。'
    for i, ch in enumerate(text_str):
        run = p.add_run(ch)
        font = run.font
        font.name = u'微软雅黑'
        # bug of python-docx
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        font.bold = (i % 2 == 0)
        font.italic = (i % 3 == 0)
        color = font.color
        color.rgb = RGBColor(i * 10 % 200 + 55, i * 20 % 200 + 55, i * 30 % 200 + 55)
    doc.save(PROJECT_PATH + '/file_location/style-2.docx')

    # 设置段落格式
    doc = Document()
    for i in range(10):
        p = doc.add_paragraph(u'段落 % d' % i)
        style = doc.styles.add_style('UserStyle%d' % i, WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.left_indent = Cm(i)
        p.style = style

    doc.save(PROJECT_PATH + '/file_location/style-3.docx')


    # 样式管理器
    doc = Document()
    for i in range(10):
        p = doc.add_paragraph(u'段落 % d' % i)
        style = doc.styles.add_style('UserStyle%d' % i, WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.left_indent = Cm(i)
        p.style = style
        if i == 7:
            style.hidden = False
            style.quick_style = True

    for style in doc.styles:
        print(style.name, style.builtin)

    doc.paragraphs[3].style = doc.styles['Subtitle']
    doc.save(PROJECT_PATH + '/file_location/style-4.docx')

if __name__ == "__main__":
    create_doc()
    insert_graph_doc()
    create_structure_doc()
    style_doc()
    read_doc_test()
    create_cell_doc_test()
    operate_cell_doc_test()
    operate_cell_doc_1_test()
    usage_of_style_doc()